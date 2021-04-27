#pip install python-docx
#start from here if python-docx is installed
import random

from docx import Document

def Masking(Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace):
    for j in range(Name_length-1):
        Name_replace=Name_replace+"X"

    for j in range(Record_length-1):
        Record_replace=Record_replace+"X"

    for j in range(Age_length-1):
        Age_replace=Age_replace+"X"

    Dictionary = {Name: " "+Name_replace, Record:" "+Record_replace, Age:" "+Age_replace}
    for i in Dictionary:
        for p in document.paragraphs:
            if p.text.find(i)>=0:
                p.text=p.text.replace(i,Dictionary[i])
    #save changed document
    document.save('Masked.docx')

def Pseudonymization(Name, Record, Age, Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace):
    #s = "Adriano Celentano"
    import random

    #for name pseudonymization
    inds_name = [i for i, _ in enumerate(Name) if not Name.isspace()]
    #length = len(Name)
    sam_name = random.sample(inds_name, Name_length-1)
    from string import ascii_letters

    lst_name = list(Name)
    # print(len(inds))
    for ind in sam_name:
        lst_name[ind] = random.choice(ascii_letters)
    Name_replace = "".join(lst_name)


    # for record pseudonymization
    inds_record = [i for i, _ in enumerate(Record) if not Record.isspace()]
    # length = len(Name)
    sam_record = random.sample(inds_record, Record_length-1)
    from string import ascii_letters

    lst_record = list(Record)
    # print(len(inds))
    for ind in sam_record:
        lst_record[ind] = random.choice(ascii_letters)
    Record_replace = "".join(lst_record)




    # for age pseudonymization
    inds_age = [i for i, _ in enumerate(Age) if not Age.isspace()]
    # length = len(Name)
    sam_age = random.sample(inds_age, Age_length-1)
    from string import ascii_letters

    lst_age = list(Age)
    # print(len(inds))
    for ind in sam_age:
        lst_age[ind] = random.choice(ascii_letters)
    Age_replace = "".join(lst_age)




    Dictionary = {Name: " " + Name_replace, Record: " " + Record_replace, Age: " " + Age_replace}
    for i in Dictionary:
        for p in document.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i, Dictionary[i])


    #print(Name_replace)
    #print(Record_replace)
    #print(Age_replace)

    # print(sam)
    # print(inds[0])
    # print(lst)
    #save changed document
    document.save('Pseudonymized.docx')








document = Document('EHR_Demo.docx')

for i in range(7):
    j=1
    if i==1:
        Name=document.paragraphs[i].runs[j].text
    elif i==2:
        Record=document.paragraphs[i].runs[j].text
    elif i==4:
        Age=document.paragraphs[i].runs[j].text

Name_length=(len(Name))
Record_length=(len(Record))
Age_length=(len(Age))

Name_replace=""
Record_replace=""
Age_replace=""

identification_list= ["Pseudonymization", "Masking"]
technique = random.choice(identification_list)
if technique=="Masking":
    Masking(Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace)
elif technique=="Pseudonymization":
    Pseudonymization(Name, Record, Age, Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace)
#Masking(Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace)
#Pseudonymization(Name, Record, Age, Name_length, Record_length, Age_length, Name_replace, Record_replace, Age_replace)








